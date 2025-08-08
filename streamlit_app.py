import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from datetime import datetime
import re
import zipfile
import logging
import html
import math

# --- Setup Logging, Constants, and Utility Functions ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

INDENT_FOR_IND_TAG_CM = 1.25
MAIN_LIST_TEXT_START_CM = 1.0  # For top-level paragraphs
SUB_LIST_TEXT_START_CM = 1.0   # Left indent for sub-paragraphs
SUB_LIST_HANGING_CM = 2.0      # Hanging indent for sub-paragraphs
SUB_ROMAN_TEXT_START_CM = 1.5  # Left indent for sub-sub-paragraphs
SUB_ROMAN_HANGING_CM = 2.5     # Hanging indent for sub-sub-paragraphs

def sanitize_input(text):
    if not isinstance(text, str): text = str(text)
    return html.escape(text)

@st.cache_data
def load_firm_details():
    return {
        "name": "Ramsdens Solicitors LLP", "short_name": "Ramsdens",
        "person_responsible_name": "Paul Pinder", "person_responsible_title": "Senior Associate",
        "supervisor_name": "Nick Armitage", "supervisor_title": "Partner",
        "person_responsible_phone": "01484 821558", "person_responsible_mobile": "07923 250815",
        "person_responsible_email": "paul.pinder@ramsdens.co.uk", "assistant_name": "Reece Collier",
        "supervisor_contact_for_complaints": "Nick Armitage on 01484 507121", "bank_name": "Barclays Bank PLC",
        "bank_address": "17 Market Place, Huddersfield", "account_name": "Ramsdens Solicitors LLP Client Account",
        "sort_code": "20-43-12", "account_number": "03909026",
        "marketing_email": "dataprotection@ramsdens.co.uk",
        "marketing_address": "Ramsdens Solicitors LLP, Oakley House, 1 Hungerford Road, Edgerton, Huddersfield, HD3 3AL"
    }

@st.cache_data
def load_precedent_text():
    try:
        with open("precedent.txt", "r", encoding="utf-8") as f:
            content = f.read()
            logger.info("Successfully loaded precedent.txt")
            return content
    except FileNotFoundError:
        st.error("precedent.txt not found. Please ensure the template text file is in the correct directory.")
        logger.error("precedent.txt not found")
        return ""
    except Exception as e:
        st.error(f"Error loading precedent.txt: {e}")
        logger.error(f"Error loading precedent.txt: {e}")
        return ""

def get_placeholder_map(app_inputs, firm_details):
    placeholders = {
        'qu1_dispute_nature': app_inputs.get('qu1_dispute_nature', ''),
        'qu2_initial_steps': app_inputs.get('qu2_initial_steps', ''),
        'qu3_timescales': app_inputs.get('qu3_timescales', ''),
        'qu4_initial_costs_with_vat': app_inputs.get('qu4_initial_costs_with_vat', 'XX,XXX'),
        'our_ref': str(app_inputs.get('our_ref', '')),
        'your_ref': str(app_inputs.get('your_ref', '')),
        'letter_date': str(app_inputs.get('letter_date', '')),
        'client_name_input': str(app_inputs.get('client_name_input', '')),
        'client_salutation': str(app_inputs.get('client_salutation', '')),
        'client_address_line1': str(app_inputs.get('client_address_line1', '')),
        'client_address_line2_conditional': str(app_inputs.get('client_address_line2_conditional', '')),
        'client_postcode': str(app_inputs.get('client_postcode', '')),
        'matter_number': str(app_inputs.get('our_ref', '')),
        'name': str(app_inputs.get('name', '')),
    }
    firm_placeholders = {k: str(v) for k, v in firm_details.items()}
    placeholders.update(firm_placeholders)
    logger.info(f"Placeholder map created: {placeholders}")
    return placeholders

def add_formatted_runs(paragraph, text_line, placeholder_map):
    processed_text = text_line
    for placeholder, value in placeholder_map.items():
        placeholder_pattern = f"{{{placeholder}}}"
        if placeholder_pattern in processed_text:
            logger.info(f"Replacing placeholder {placeholder_pattern} with {value}")
        processed_text = processed_text.replace(placeholder_pattern, str(value))
    parts = re.split(r'(<bd>|</bd>|<ins>|</ins>)', processed_text)
    is_bold = is_underline = False
    for part in parts:
        if not part: continue
        if part == "<bd>": is_bold = True
        elif part == "</bd>": is_bold = False
        elif part == "<ins>": is_underline = True
        elif part == "</ins>": is_underline = False
        else:
            for i, line_part in enumerate(part.split('\n')):
                if i > 0: paragraph.add_run().add_break()
                run = paragraph.add_run(line_part)
                run.bold, run.underline = is_bold, is_underline
                run.font.name, run.font.size = 'Arial', Pt(11)
                logger.debug(f"Added run: {line_part}, bold={is_bold}, underline={is_underline}")

def should_render_track_block(tag, claim_assigned, selected_track):
    tag_map = {
        'a1': (True, "Small Claims Track"), 'a2': (True, "Fast Track"),
        'a3': (True, "Intermediate Track"), 'a4': (True, "Multi Track"),
        'u1': (False, "Small Claims Track"), 'u2': (False, "Fast Track"),
        'u3': (False, "Intermediate Track"), 'u4': (False, "Multi Track")
    }
    expected = tag_map.get(tag)
    if not expected: return False
    return claim_assigned == expected[0] and selected_track == expected[1]

def generate_initial_advice_doc(app_inputs, placeholder_map):
    doc = Document()
    doc.styles['Normal'].font.name, doc.styles['Normal'].font.size = 'Arial', Pt(11)
    p = doc.add_paragraph()
    add_formatted_runs(p, "Initial Advice Summary - Matter Number: {matter_number}", placeholder_map)
    p.paragraph_format.space_after = Pt(12)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    rows_data = [
        ("Date of Advice", app_inputs['initial_advice_date'].strftime('%d/%m/%Y') if app_inputs.get('initial_advice_date') else ''),
        ("Method of Advice", app_inputs.get('initial_advice_method', '')),
        ("Advice Given", app_inputs.get('initial_advice_content', ''))
    ]
    for i, (label, value) in enumerate(rows_data):
        table.rows[i].cells[0].text, table.rows[i].cells[1].text = label, value
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def preprocess_precedent(precedent_content, app_inputs):
    logical_elements = []
    lines = precedent_content.splitlines()
    i = 0
    current_block_tag = None
    block_lines = []
    current_list_type = None  # Track list type: 'numbered', 'letter', 'roman'

    def determine_list_type(line):
        line = line.strip()
        if line.startswith('<a>'):
            return 'letter'
        elif line.startswith('<i>'):
            return 'roman'
        elif line.startswith('1.'):
            return 'numbered'
        return None

    def flush_block(block_tag, block_lines, list_type):
        if not block_lines:
            return
        content = block_lines[0].strip() if block_lines else ""
        if not content:
            logical_elements.append({'type': 'blank_line', 'content_lines': [], 'block_tag': block_tag, 'list_type': None})
        elif '<ins>' in content:
            logical_elements.append({'type': 'heading', 'content_lines': block_lines, 'block_tag': block_tag, 'list_type': None})
        elif '[FEE_TABLE_PLACEHOLDER]' in content:
            logical_elements.append({'type': 'fee_table', 'content_lines': block_lines, 'block_tag': block_tag, 'list_type': None})
        else:
            element_type = list_type or 'general_paragraph'
            logical_elements.append({
                'type': element_type,
                'content_lines': block_lines,
                'block_tag': block_tag,
                'list_type': list_type
            })

    while i < len(lines):
        line = lines[i].strip()
        match_start_tag = re.match(r'^\[(indiv|corp|a[1-4]|u[1-4])\]$', line)
        match_end_tag = re.match(r'^\[/(indiv|corp|a[1-4]|u[1-4])\]$', line)
        if match_start_tag:
            if block_lines and current_block_tag is None:
                flush_block(None, block_lines, current_list_type)
                block_lines = []
            current_block_tag = match_start_tag.group(1)
            i += 1
            continue
        elif match_end_tag:
            if block_lines and current_block_tag == match_end_tag.group(1):
                flush_block(current_block_tag, block_lines, current_list_type)
                block_lines = []
            current_block_tag = None
            current_list_type = None
            i += 1
            continue
        else:
            if line:
                new_list_type = determine_list_type(line)
                if new_list_type and new_list_type != current_list_type:
                    if block_lines:
                        flush_block(current_block_tag, block_lines, current_list_type)
                        block_lines = []
                    current_list_type = new_list_type
                block_lines.append(lines[i])
            else:
                if block_lines:
                    flush_block(current_block_tag, block_lines, current_list_type)
                    block_lines = []
                logical_elements.append({'type': 'blank_line', 'content_lines': [], 'block_tag': current_block_tag, 'list_type': None})
            i += 1

    if block_lines:
        flush_block(current_block_tag, block_lines, current_list_type)

    return logical_elements

def process_precedent_text(precedent_content, app_inputs, placeholder_map):
    try:
        doc = Document()
        doc.styles['Normal'].font.name, doc.styles['Normal'].font.size = 'Arial', Pt(11)
        numbering_elm = doc.part.numbering_part.element
        abstract_num_id, num_instance_id = 10, 1

        abstract_num = OxmlElement('w:abstractNum')
        abstract_num.set(qn('w:abstractNumId'), str(abstract_num_id))

        def create_level(ilvl, numFmt, lvlText, left_indent_cm, hanging_cm, start_val=1):
            lvl = OxmlElement('w:lvl')
            lvl.set(qn('w:ilvl'), str(ilvl))
            numFmt_el = OxmlElement('w:numFmt')
            numFmt_el.set(qn('w:val'), numFmt)
            lvl.append(numFmt_el)
            lvlText_el = OxmlElement('w:lvlText')
            lvlText_el.set(qn('w:val'), lvlText)
            lvl.append(lvlText_el)
            start_el = OxmlElement('w:start')
            start_el.set(qn('w:val'), str(start_val))
            lvl.append(start_el)
            pPr = OxmlElement('w:pPr')
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), str(Cm(left_indent_cm).twips))
            ind.set(qn('w:hanging'), str(Cm(hanging_cm).twips))
            pPr.append(ind)
            lvl.append(pPr)
            return lvl

        # Level 0: Numbered list (1.), left margin (0cm), text at 1cm (hanging 1cm)
        abstract_num.append(create_level(0, 'decimal', '%1.', 0, 1.0, start_val=1))
        # Level 1: Letter list (a), left at 1cm, text at 2cm (hanging 1cm)
        abstract_num.append(create_level(1, 'lowerLetter', '(%2)', SUB_LIST_TEXT_START_CM, 1.0, start_val=1))
        # Level 2: Roman list (i), left at 1.5cm, text at 2.5cm (hanging 1cm)
        abstract_num.append(create_level(2, 'lowerRoman', '(%3)', SUB_ROMAN_TEXT_START_CM, 1.0, start_val=1))
        numbering_elm.append(abstract_num)

        num = OxmlElement('w:num')
        num.set(qn('w:numId'), str(num_instance_id))
        abstract_num_id_ref = OxmlElement('w:abstractNumId')
        abstract_num_id_ref.set(qn('w:val'), str(abstract_num_id))
        num.append(abstract_num_id_ref)
        numbering_elm.append(num)

        logical_elements = preprocess_precedent(precedent_content, app_inputs)
        for element in logical_elements:
            render_this_element = True
            tag = element.get('block_tag')
            if tag:
                if tag == 'indiv':
                    render_this_element = app_inputs['client_type'] == 'Individual'
                elif tag == 'corp':
                    render_this_element = app_inputs['client_type'] == 'Corporate'
                else:
                    render_this_element = should_render_track_block(tag, app_inputs['claim_assigned'], app_inputs['selected_track'])
            if not render_this_element:
                continue

            content = element['content_lines'][0] if element['content_lines'] else ""
            logger.debug(f"Processing element: {element['type']}, content: {content}")

            def add_list_item(level, text):
                p = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                numPr = pPr.get_or_add_numPr()
                numPr.get_or_add_ilvl().val = level
                numPr.get_or_add_numId().val = num_instance_id
                cleaned_content = text.replace('<a>', '').replace('<i>', '').strip()
                add_formatted_runs(p, cleaned_content, placeholder_map)
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p.paragraph_format.space_after = Pt(6)
                logger.info(f"Added list item at level {level}: {cleaned_content}")

            if element['type'] == 'blank_line':
                continue
            elif element['type'] == 'fee_table':
                table = doc.add_table(rows=5, cols=2)
                table.style = 'Table Grid'
                fee_data = [
                    ("Grade A", "£450 (Partners, Solicitors over 8 years)"),
                    ("Grade B", "£350 (Solicitors/Legal Executives over 4 years)"),
                    ("Grade C", f"£{app_inputs['hourly_rate']} (Solicitors/Legal Executives under 4 years)"),
                    ("Grade D", "£250 (Trainees, Paralegals)"),
                    ("Grade E", "£150 (Support Staff)")
                ]
                for i, (grade, rate) in enumerate(fee_data):
                    table.rows[i].cells[0].text = grade
                    table.rows[i].cells[1].text = rate
                doc.add_paragraph().paragraph_format.space_after = Pt(12)
            elif element['type'] == 'heading':
                p = doc.add_paragraph()
                add_formatted_runs(p, content, placeholder_map)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
            elif element['type'] == 'numbered':
                add_list_item(0, content)
            elif element['type'] == 'letter':
                add_list_item(1, content)
            elif element['type'] == 'roman':
                add_list_item(2, content)
            elif element['type'] == 'general_paragraph':
                p = doc.add_paragraph()
                cleaned_content = content.replace('[ind]', '').strip()
                if '[ind]' in content:
                    p.paragraph_format.left_indent = Cm(INDENT_FOR_IND_TAG_CM)
                add_formatted_runs(p, cleaned_content, placeholder_map)
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p.paragraph_format.space_after = Pt(12)
        return doc
    except Exception as e:
        logger.error(f"Error processing precedent text: {e}", exc_info=True)
        raise

st.set_page_config(layout="wide", page_title="Ramsdens Client Care Letter Generator")
st.title("Ramsdens Client Care Letter Generator")

firm_details = load_firm_details()
precedent_content = load_precedent_text()
if not precedent_content:
    st.stop()

with st.form("input_form"):
    st.header("1. Letter & Client Details")
    c1, c2 = st.columns(2)
    with c1:
        our_ref = st.text_input("Our Reference", "PDP/10011/001")
        your_ref = st.text_input("Your Reference", "REF")
        letter_date = st.date_input("Letter Date", datetime.today())
    with c2:
        client_name_input = st.text_input("Client Full Name / Company Name", "Mr. John Smith")
        client_salutation_name = st.text_input("Salutation (for 'Dear ...' and address block)", "Mr. Smith")
        client_address_line1 = st.text_input("Address Line 1", "123 Example Street")
        client_address_line2 = st.text_input("Address Line 2 (optional)", "SomeTown")
        client_postcode = st.text_input("Postcode", "EX4 MPL")
        client_type = st.radio("Client Type", ("Individual", "Corporate"), horizontal=True)

    st.header("2. Initial Advice & Case Details")
    c1, c2 = st.columns(2)
    with c1:
        st.subheader("Initial Advice Summary")
        initial_advice_content = st.text_area("Advice Given", "Advised on merits...", height=100)
        initial_advice_method = st.selectbox("Method", ["Phone Call", "In Person", "Teams Call"])
        initial_advice_date = st.date_input("Date", datetime.today())
    with c2:
        st.subheader("Case Track")
        claim_assigned_input = st.radio("Is claim already assigned?", ("Yes", "No"), horizontal=True, index=1)
        selected_track = st.selectbox("Which track applies?", ["Small Claims Track", "Fast Track", "Intermediate Track", "Multi Track"])

    st.header("3. Dynamic Content")
    qu1_dispute_nature = st.text_area('Dispute Nature', "a contractual matter...", height=75)
    qu2_initial_steps = st.text_area('Initial Work', "review documentation...", height=100)
    qu3_timescales = st.text_area("Estimated Timescales", "approx two to four weeks...", height=100)
    
    st.subheader("Estimated Initial Costs")
    hourly_rate = st.number_input("Your Hourly Rate (£)", 295)
    step_value = float(hourly_rate / 2)  # Half-hour increments
    lower_hours = 2.5
    upper_hours = 3.5
    cost_type_is_range = st.toggle("Use a cost range", True)
    if cost_type_is_range:
        lower_cost = st.number_input("Lower £", value=hourly_rate * lower_hours, step=step_value)
        upper_cost = st.number_input("Upper £", value=hourly_rate * upper_hours, step=step_value)
    else:
        fixed_cost = st.number_input("Fixed cost £", value=hourly_rate * lower_hours, step=step_value)

    submitted = st.form_submit_button("Generate Documents")

if submitted:
    if cost_type_is_range:
        lower_cost_vat = lower_cost * 1.2
        upper_cost_vat = upper_cost * 1.2
        lower_cost_vat_rounded = math.ceil(lower_cost_vat / 50) * 50
        upper_cost_vat_rounded = math.ceil(upper_cost_vat / 50) * 50
        costs_text = f"from £{lower_cost:,.2f} to £{upper_cost:,.2f} plus VAT which means between £{lower_cost_vat_rounded:,.2f} to £{upper_cost_vat_rounded:,.2f}"
    else:
        fixed_cost_vat = fixed_cost * 1.2
        fixed_cost_vat_rounded = math.ceil(fixed_cost_vat / 50) * 50
        costs_text = f"a fixed fee of £{fixed_cost:,.2f} plus VAT that being £{fixed_cost_vat_rounded:,.2f}"
    app_inputs = {
        'qu1_dispute_nature': sanitize_input(qu1_dispute_nature),
        'qu2_initial_steps': sanitize_input(qu2_initial_steps),
        'qu3_timescales': sanitize_input(qu3_timescales),
        'qu4_initial_costs_with_vat': costs_text,
        'client_type': client_type,
        'claim_assigned': claim_assigned_input == "Yes",
        'selected_track': selected_track,
        'our_ref': sanitize_input(our_ref),
        'your_ref': sanitize_input(your_ref),
        'letter_date': letter_date.strftime('%d %B %Y'),
        'client_name_input': sanitize_input(client_name_input),
        'client_salutation': sanitize_input(client_salutation_name),
        'client_address_line1': sanitize_input(client_address_line1),
        'client_address_line2_conditional': sanitize_input(client_address_line2) if client_address_line2 else "",
        'client_postcode': sanitize_input(client_postcode),
        'name': sanitize_input(firm_details["person_responsible_name"]),
        'initial_advice_content': initial_advice_content,
        'initial_advice_method': initial_advice_method,
        'initial_advice_date': initial_advice_date,
        'hourly_rate': hourly_rate,
        'firm_details': firm_details
    }
    placeholder_map = get_placeholder_map(app_inputs, firm_details)
    try:
        doc = process_precedent_text(precedent_content, app_inputs, placeholder_map)
        client_care_doc_io = io.BytesIO()
        doc.save(client_care_doc_io)
        client_care_doc_io.seek(0)
        
        advice_doc_io = generate_initial_advice_doc(app_inputs, placeholder_map)
        
        client_name_safe = re.sub(r'[^\w\s-]', '', client_name_input).strip().replace(' ', '_')
        zip_io = io.BytesIO()
        
        with zipfile.ZipFile(zip_io, 'w', zipfile.ZIP_DEFLATED) as zipf:
            zipf.writestr(f"Client_Care_Letter_{client_name_safe}.docx", client_care_doc_io.getvalue())
            if advice_doc_io:
                zipf.writestr(f"Initial_Advice_Summary_{client_name_safe}.docx", advice_doc_io.getvalue())
        
        zip_io.seek(0)
        
        st.success("Documents Generated Successfully!")
        st.download_button("Download All Documents as ZIP", zip_io, f"Client_Docs_{client_name_safe}.zip", "application/zip")
        
    except Exception as e:
        st.error(f"An error occurred while building the documents: {e}")
        logger.exception("Error during document generation:")
